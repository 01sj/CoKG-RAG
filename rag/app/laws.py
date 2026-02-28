#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.



#  

import logging
import re
from io import BytesIO
from docx import Document

from deepdoc.parser.utils import get_text
from rag.nlp import bullets_category, remove_contents_table, \
    make_colon_as_title, tokenize_chunks, docx_question_level, tree_merge
from rag.nlp import rag_tokenizer, Node




class Docx:
    def __init__(self):
        pass

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def old_call(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page and p.text.strip():
                lines.append(self.__clean(p.text))
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        return [line for line in lines if line]

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
            self.doc = Document(
                filename) if not binary else Document(BytesIO(binary))
            pn = 0
            lines = []
            level_set = set()
            bull = bullets_category([p.text for p in self.doc.paragraphs])
            
            # 从文件名中提取法律名称
            # 文件名格式通常是：《法律名称》，第XX号法令_format.docx
            # 或者：《法律名称》.docx
            law_name = None
            if filename:
                import os
                base_name = os.path.basename(filename)
                # 尝试匹配《...》格式的法律名称
                match = re.search(r'《([^》]+)》', base_name)
                if match:
                    law_name = f"《{match.group(1)}》"
                else:
                    # 如果没有找到，尝试从文件名的开头提取（去除扩展名）
                    name_without_ext = os.path.splitext(base_name)[0]
                    # 去除可能的格式后缀（如_format）
                    name_without_ext = re.sub(r'_format$', '', name_without_ext)
                    # 提取到第一个逗号或空格之前的内容
                    match = re.match(r'^([^，,]+)', name_without_ext)
                    if match:
                        law_name = match.group(1).strip()
            
            # 定义各层级的匹配模式（需要在第一步之前定义，用于拆分复合段落）
            structure_patterns = {
                '编': r'^[\s\u3000]*第[零一二三四五六七八九十百0-9]+(分?编|部分)',
                '章': r'^[\s\u3000]*第[零一二三四五六七八九十百0-9]+章',
                '节': r'^[\s\u3000]*第[零一二三四五六七八九十百0-9]+节',
                '条': r'^[\s\u3000]*第[零一二三四五六七八九十百0-9]+条',
            }
            
            # 第一步：收集所有段落的级别和文本
            # 对于包含换行符的复合段落，只拆分出结构标题（编/章/节/条），其余内容保持在对应的条目中
            for p in self.doc.paragraphs:
                if pn > to_page:
                    break
                question_level, p_text = docx_question_level(p, bull)
                if not p_text.strip("\n"):
                    continue

                # 如果段落包含换行符，检查是否包含结构标题
                if '\n' in p_text and len(p_text.split('\n')) > 1:
                    sub_lines = p_text.split('\n')
                    buffered_content = []
                    current_struct_level = question_level

                    def flush_buffer():
                        nonlocal buffered_content, current_struct_level
                        if buffered_content:
                            content_text = '\n'.join(buffered_content)
                            # 使用比结构级别更深的级别，确保在build_tree中被合并到当前节点
                            content_level = current_struct_level + 1
                            lines.append((content_level, content_text))
                            level_set.add(content_level)
                            buffered_content = []

                    for sub_line in sub_lines:
                        sub_line = sub_line.strip()
                        if not sub_line:
                            continue

                        matched_structure = False
                        for struct_name, pattern in structure_patterns.items():
                            if re.match(pattern, sub_line):
                                flush_buffer()
                                struct_level_map = {'编': 1, '章': 2, '节': 3, '条': 4}
                                current_struct_level = struct_level_map.get(struct_name, question_level)
                                lines.append((current_struct_level, sub_line))
                                level_set.add(current_struct_level)
                                matched_structure = True
                                break

                        if not matched_structure:
                            buffered_content.append(sub_line)

                    flush_buffer()
                else:
                    # 普通段落，直接添加
                    lines.append((question_level, p_text))
                    level_set.add(question_level)
                
                for run in p.runs:
                    if 'lastRenderedPageBreak' in run._element.xml:
                        pn += 1
                        continue
                    if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                        pn += 1

            sorted_levels = sorted(level_set)
            
            # 第二步：智能识别文档层级结构
            # BULLET_PATTERN[0]定义：0=编, 1=章, 2=节, 3=条, 4=括号标题
            # 策略：识别编、章、节、条的实际级别，而不是依赖固定级别
            # 注意：structure_patterns已在第一步之前定义
            
            # 统计每个层级在各个级别中出现的次数
            structure_level_count = {name: {} for name in structure_patterns.keys()}
            structure_samples = {name: {} for name in structure_patterns.keys()}
            
            for lvl, txt in lines:
                # 注意：txt可能已经经过strip处理，但正则表达式中的 ^ ? 可以匹配可选的单个空格
                # 为了兼容标题前有一个空格的情况，我们直接使用txt进行匹配
                # 正则表达式 ^ ? 表示：行首可选的一个空格
                for struct_name, pattern in structure_patterns.items():
                    if re.match(pattern, txt):
                        structure_level_count[struct_name][lvl] = structure_level_count[struct_name].get(lvl, 0) + 1
                        if lvl not in structure_samples[struct_name]:
                            structure_samples[struct_name][lvl] = []
                        if len(structure_samples[struct_name][lvl]) < 2:
                            # 保存样本时去除前导空格，保持格式一致
                            structure_samples[struct_name][lvl].append(txt.lstrip()[:20])
            
            # 识别每个层级的实际级别（选择数量最多的级别）
            structure_levels = {}
            for struct_name, level_counts in structure_level_count.items():
                if level_counts:
                    actual_level = max(level_counts.keys(), key=lambda k: level_counts[k])
                    structure_levels[struct_name] = {
                        'level': actual_level,
                        'count': level_counts[actual_level],
                        'samples': structure_samples[struct_name].get(actual_level, [])
                    }
            
            # 第三步：确定目标深度
            print(f"\n{'='*60}")
            print(f"文档结构分析：")
            print(f"  - 所有级别: {sorted_levels}")
            
            # 显示识别到的文档结构
            if structure_levels:
                print(f"  - 文档层级结构:")
                for struct_name in ['编', '章', '节', '条']:
                    if struct_name in structure_levels:
                        info = structure_levels[struct_name]
                        print(f"    * {struct_name}: 级别{info['level']}, 数量{info['count']}, 样本{info['samples']}")
            
            # 如果识别到"条"，使用"条"的级别来确定depth
            if '条' in structure_levels:
                article_level = structure_levels['条']['level']
                article_count = structure_levels['条']['count']
                
                # 关键修正：depth应该设置为"条"的级别本身，而不是父级别
                # 这样在build_tree中：
                # - level <= depth（包括"条"）：创建节点
                # - level > depth（条的内容）：添加到当前节点
                # 在_dfs中：
                # - level == article_level的节点会成为独立块
                target_depth = article_level
                
                # 找到"条"的父级别用于说明
                hierarchy = ['编', '章', '节']
                parent_names = []
                for parent_struct in ['编', '章', '节']:
                    if parent_struct in structure_levels and structure_levels[parent_struct]['level'] < article_level:
                        parent_names.append(f"'{parent_struct}'")
                
                if parent_names:
                    print(f"  - 分块策略: {'/'.join(parent_names)}作为标题路径")
                else:
                    print(f"  - 分块策略: 直接按'条'分块（无上级层级）")
                
                print(f"  - '条'(级别{article_level})作为分块单元")
                print(f"  - '条'的内容(级别>{article_level})合并到对应的'条'中")
                print(f"  - 设置 depth={target_depth}")
                print(f"  - 预计生成约 {article_count} 个分块")
            else:
                # 如果没找到"条"，使用默认策略
                print(f"  - 警告：未识别到'条'，使用默认分块策略")
                if len(sorted_levels) > 1:
                    target_depth = sorted_levels[-2]
                else:
                    target_depth = sorted_levels[0] if sorted_levels else 1
                print(f"  - 设置 depth={target_depth}")
            
            print(f"{'='*60}\n")
            
            # 第四步：构建树并生成分块
            root = Node(level=0, depth=target_depth, texts=[], law_name=law_name)
            root.build_tree(lines)

            return [element for element in root.get_tree() if element]


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
        Supported file formats are docx, txt.
    """
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": 512, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    sections = []
    # is it English
    eng = lang.lower() == "english"  # is_english(sections)

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunks = Docx()(filename, binary)
        callback(0.7, "Finish parsing.")
        return tokenize_chunks(chunks, doc, eng, None)
    
    elif re.search(r"\.(txt|md|markdown|mdx)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        sections = txt.split("\n")
        sections = [s for s in sections if s]
        callback(0.8, "Finish parsing.")
    else:
        raise NotImplementedError(
            "file type not supported yet. Supported formats: docx, txt, md, markdown")

    # Remove 'Contents' part
    remove_contents_table(sections, eng)

    make_colon_as_title(sections)
    bull = bullets_category(sections)
    res = tree_merge(bull, sections, 2)

    if not res:
        callback(0.99, "No chunk parsed out.")

    return tokenize_chunks(res, doc, eng, None)

    # chunks = hierarchical_merge(bull, sections, 5)
    #     return tokenize_chunks(["\n".join(ck)for ck in chunks], doc, eng, pdf_parser)

if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass
    chunk(sys.argv[1], callback=dummy)
